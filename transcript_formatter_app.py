import streamlit as st
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.shared import qn, OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import tempfile

st.set_page_config(page_title="Transcript Formatter", layout="centered")
st.title("📄 Coaching Transcript Formatter")

uploaded_file = st.file_uploader("Upload your .vtt transcript file", type=["vtt"])

def add_page_number_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")

    # PAGE field
    run_page = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.text = "PAGE"

    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')

    fld_text = OxmlElement('w:t')
    fld_text.text = "1"

    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')

    run_page._r.extend([fld_char1, instr_text, fld_char2, fld_text, fld_char3])

    paragraph.add_run(" of ")

    # NUMPAGES field
    run_total = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.text = "NUMPAGES"

    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')

    fld_text = OxmlElement('w:t')
    fld_text.text = "1"

    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')

    run_total._r.extend([fld_char1, instr_text, fld_char2, fld_text, fld_char3])

if uploaded_file:
    content = uploaded_file.read().decode("utf-8")
    lines = content.splitlines()

    # Extract speaker names
    speaker_set = set()
    for line in lines:
        match = re.match(r"(.+?):\s+.*", line.strip())
        if match:
            speaker_set.add(match.group(1).strip())
        if len(speaker_set) == 2:
            break

    if len(speaker_set) != 2:
        st.error("Could not find exactly two speakers in the file.")
        st.stop()

    speaker_list = sorted(speaker_set)
    coach_display_name = st.radio("Who is the Coach?", speaker_list)

    coach_name = coach_display_name.lower()
    client_name = [s for s in speaker_list if s != coach_display_name][0]

    # Parse transcript entries
    entries = []
    entry_num = 1
    current_speaker = None
    current_text = ""
    current_timestamp = ""

    for line in lines:
        line = line.strip()
        timestamp_match = re.match(r"(\d{2}:\d{2}:\d{2}\.\d{3}) --> \d{2}:\d{2}:\d{2}\.\d{3}", line)
        if timestamp_match:
            current_timestamp = timestamp_match.group(1)
            continue
        speaker_match = re.match(r"(.+?):\s*(.*)", line)
        if speaker_match:
            speaker_name = speaker_match.group(1).strip()
            text = speaker_match.group(2)
            if current_speaker and current_text:
                entries.append((entry_num, current_timestamp, current_speaker, current_text))
                entry_num += 1
            current_speaker = speaker_name
            current_text = text
        elif line and current_speaker:
            current_text += " " + line
        elif not line:
            if current_speaker and current_text:
                entries.append((entry_num, current_timestamp, current_speaker, current_text))
                entry_num += 1
                current_speaker = None
                current_text = ""

    if current_speaker and current_text:
        entries.append((entry_num, current_timestamp, current_speaker, current_text))

    # Group by continuity
    grouped_entries = []
    start_num = entries[0][0]
    start_time = entries[0][1]
    current_speaker = entries[0][2]
    aggregated_text = entries[0][3]
    end_num = start_num

    for i in range(1, len(entries)):
        num, time, speaker, text = entries[i]
        if speaker == current_speaker:
            aggregated_text += " " + text
            end_num = num
        else:
            grouped_entries.append((start_num, end_num, start_time, current_speaker, aggregated_text))
            start_num = num
            end_num = num
            start_time = time
            current_speaker = speaker
            aggregated_text = text
    grouped_entries.append((start_num, end_num, start_time, current_speaker, aggregated_text))

    # Build document
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Header
    header_para = section.header.paragraphs[0]
    header_run = header_para.add_run(coach_display_name)
    header_run.font.color.rgb = RGBColor(64, 64, 64)
    header_run.font.size = Pt(14)
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Footer
    footer_para = section.footer.paragraphs[0]
    add_page_number_field(footer_para)

    # Title
    doc.add_heading("Coaching Session Transcript with Feedback", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Legend
    legend_items = [
        ("SD", "Evidence of competency demonstration. Demonstrating the skill at one point does not mean the skill was demonstrated throughout the session."),
        ("LD", "Lack of evidence of demonstration or contra evidence of competency in the marked moment of the discussion."),
        ("AMDOS", "Ask Me During Our Session"),
        ("SWMDOS", "Share With Me During Our Session"),
        ("CEQ", "Close Ended Question"),
        ("ECNN", "Expansive conversation not needed."),
        ("CD", "Cognitive Distortion")
    ]
    for label, desc in legend_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = 0
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)

    for _ in range(2):
        doc.add_paragraph("")

    # Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].repeat_on_every_page = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Coaching Transcript'
    hdr_cells[1].text = "Mentor's Feedback"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for start, end, time, speaker, text in grouped_entries:
        row_cells = table.add_row().cells
        para = row_cells[0].paragraphs[0]
        range_str = f"{start}" if start == end else f"{start}-{end}"
        para.add_run(f"{range_str} [")
        time_run = para.add_run(time)
        time_run.font.color.rgb = RGBColor(105, 105, 105)
        para.add_run("] ")
        role = "Coach" if speaker.strip().lower() == coach_name else "Client"
        speaker_run = para.add_run(f"{role} {speaker}")
        speaker_run.bold = True
        para.add_run(f" {text}")
        row_cells[1].text = ""

    doc.add_paragraph("")
    doc.add_paragraph("Strengths:")
    doc.add_paragraph("")
    doc.add_paragraph("Progression Ideas:")

    # Save to temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with open(tmp_path, "rb") as f:
        st.download_button("📥 Download Formatted Transcript", f, file_name="Formatted_Transcript.docx")
